"""Document loading for hybrid RAG.

Supported formats: PDF, DOCX, TXT, Markdown.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, List, Optional


@dataclass
class LoadedDocument:
    text: str
    metadata: dict


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}

    def __init__(self, base_timestamp: Optional[str] = None):
        self.base_timestamp = base_timestamp

    def load_paths(self, paths: Iterable[Path]) -> List[LoadedDocument]:
        docs: List[LoadedDocument] = []
        for path in paths:
            docs.extend(self.load_path(Path(path)))
        return docs

    def load_path(self, path: Path) -> List[LoadedDocument]:
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported document format: {ext}")
        if ext == ".pdf":
            return self._load_pdf(path)
        if ext == ".docx":
            return self._load_docx(path)
        if ext in {".md", ".markdown"}:
            return self._load_markdown(path)
        return self._load_text(path)

    def _now(self) -> str:
        if self.base_timestamp:
            return self.base_timestamp
        return datetime.now(timezone.utc).isoformat()

    def _base_metadata(self, path: Path) -> dict:
        return {
            "document": path.name,
            "page": 1,
            "section": "default",
            "timestamp": self._now(),
        }

    def _load_text(self, path: Path) -> List[LoadedDocument]:
        text = path.read_text(encoding="utf-8", errors="replace")
        return [LoadedDocument(text=text, metadata=self._base_metadata(path))]

    def _load_markdown(self, path: Path) -> List[LoadedDocument]:
        text = path.read_text(encoding="utf-8", errors="replace")
        lines = text.splitlines()

        docs: List[LoadedDocument] = []
        current_section = "default"
        current_lines: List[str] = []

        def flush_section() -> None:
            if not current_lines:
                return
            body = "\n".join(current_lines).strip()
            if not body:
                return
            md = self._base_metadata(path)
            md["section"] = current_section
            docs.append(LoadedDocument(text=body, metadata=md))

        for line in lines:
            if line.lstrip().startswith("#"):
                flush_section()
                current_lines = []
                current_section = line.lstrip("#").strip() or "default"
            else:
                current_lines.append(line)
        flush_section()

        if not docs:
            docs.append(LoadedDocument(text=text, metadata=self._base_metadata(path)))
        return docs

    def _load_pdf(self, path: Path) -> List[LoadedDocument]:
        try:
            from pypdf import PdfReader
        except Exception as exc:
            raise ImportError("pypdf is required for PDF loading.") from exc

        reader = PdfReader(str(path))
        docs: List[LoadedDocument] = []
        for i, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            md = self._base_metadata(path)
            md["page"] = i
            md["section"] = f"page_{i}"
            docs.append(LoadedDocument(text=text, metadata=md))
        return docs

    def _load_docx(self, path: Path) -> List[LoadedDocument]:
        try:
            from docx import Document
        except Exception as exc:
            raise ImportError("python-docx is required for DOCX loading.") from exc

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        text = "\n".join(paragraphs).strip()
        if not text:
            return []
        md = self._base_metadata(path)
        return [LoadedDocument(text=text, metadata=md)]
